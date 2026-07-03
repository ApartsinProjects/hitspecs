# -*- coding: utf-8 -*-
"""Build the HIT catalogue package for CY5: Blockchain & Decentralized Systems.

The script clones the existing HIT DOCX templates from the AI Systems course
package and replaces the body content while preserving template headers,
styles, and page setup.
"""
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_ROOT = r"E:\Projects\Courses\AIApsCourse\build\templates"
COMMITTEE_TEMPLATE = r"E:\Projects\Courses\AIApsCourse\hit-catalogue\committee_questionnaire.docx"
OUT = os.path.join(ROOT, "hit-catalogue", "CY5")
os.makedirs(OUT, exist_ok=True)

TITLE_EN = "Blockchain & Decentralized Systems"
TITLE_HE = "בלוקצ'יין ומערכות מבוזרות"
OFFICIAL_TITLE_HE = "בלוקצ'יין חזון ופרקטיקה"
COURSE_NO = "67007"
LECTURER_HE = "מיכה ברשפ"
LECTURER_EN = "Micha Barshap"

HOURS_HE = "שעות שבועיות: הרצאה 3 שעות + תרגול 1 שעה, סה\"כ שעות - 4"
HOURS_EN = "Weekly hours: lecture 3 hours + practice 1 hour, total 4 hours"
CREDITS_HE = "נקודות זכות: 3.5"
CREDITS_EN = "3.5 credits"
PREREQ_HE = "דרישות קדם: מבני נתונים 61104, תכנות מונחה עצמים 61307"
PREREQ_EN = "Prerequisites: Data Structures 61104, Object-Oriented Programming 61307"
CONCURRENT_HE = "דרישות במקביל: אין"
CONCURRENT_EN = "Concurrent requirements: none"

BIB = [
    "Nakamoto, Satoshi. Bitcoin: A Peer-to-Peer Electronic Cash System. 2008.",
    "Antonopoulos, Andreas M., and David A. Harding. Mastering Bitcoin, 3rd Edition. O'Reilly Media, 2023.",
    "Antonopoulos, Andreas M., and Gavin Wood. Mastering Ethereum. O'Reilly Media, 2018.",
    "Narayanan, Arvind, Joseph Bonneau, Edward Felten, Andrew Miller, and Steven Goldfeder. Bitcoin and Cryptocurrency Technologies. Princeton University Press, 2016.",
    "Lamport, Leslie, Robert Shostak, and Marshall Pease. The Byzantine Generals Problem. ACM Transactions on Programming Languages and Systems, 1982.",
    "Castro, Miguel, and Barbara Liskov. Practical Byzantine Fault Tolerance. OSDI, 1999.",
    "Garay, Juan, Aggelos Kiayias, and Nikos Leonardos. The Bitcoin Backbone Protocol: Analysis and Applications. EUROCRYPT, 2015.",
    "Ethereum Foundation. Solidity Documentation.",
]

WEEKS_EN = [
    ("1", "Blockchain foundations: blocks, transactions, miners, wallets, UTXOs, mempools, forks, hash functions, public and private keys, and 51 percent attacks."),
    ("2", "Peer-to-peer networking and data structures: client-server and P2P communication, replicated ledgers, message ordering, and failure modes."),
    ("3", "Blockchain data structures: Merkle trees, Bloom filters, transaction validation, signed transactions, block integrity, and fork-aware acceptance."),
    ("4", "Consensus: Byzantine generals, PBFT, proof of work, proof of stake, Nakamoto consensus, longest-chain selection, and finality."),
    ("5", "Design specification and threat model: privacy-preserving systems, hard and soft forks, layer-2 scaling, Lightning/Plasma ideas, and system trade-offs."),
    ("6", "Ethereum and smart contracts: accounts, EVM, gas, Solidity, local chains, Remix, Hardhat or Foundry, Ganache, MetaMask, ethers.js and Web3.js."),
    ("7", "Ethereum ecosystem: token economics, DAOs, ICO-style crowdfunding, governance, and social and economic implications."),
    ("8", "Secure DApp development: DApp client implementation, smart-contract testing, OpenZeppelin libraries, reentrancy, access control, and oracle manipulation."),
    ("9", "Tokens: ERC-20, ERC-721, token balances, minting, transfers, approvals, NFT marketplace patterns, and audit considerations."),
    ("10", "Scaling and storage: Ethereum 2.0 concepts, proof-of-stake finality, IPFS content addressing, DApp metadata, and alternative networks."),
    ("11", "Wallets and custody: HD wallets, seed phrases, hot and cold storage, multi-signature approval, recovery, and revocation scenarios."),
    ("12", "DeFi: AMMs, lending, liquidity, tokenomics, flash-loan attacks, oracle attacks, tracing, auditing, and hardening."),
    ("13", "Final project and trends: working decentralized system, audit report, oral defense, and current blockchain directions."),
]

WEEKS_HE = [
    ("1", "יסודות בלוקצ'יין: בלוקים, עסקאות, כורים, ארנקים, UTXO, mempool, fork, פונקציות גיבוב, מפתח פרטי וציבורי, והתקפת 51 אחוז."),
    ("2", "רשתות עמית לעמית ומבני נתונים: תקשורת שרת-לקוח ו-P2P, ספרי חשבונות משוכפלים, סדר הודעות ומודלי כשל."),
    ("3", "מבני נתונים בבלוקצ'יין: עצי Merkle, Bloom filters, אימות עסקאות, עסקאות חתומות, שלמות בלוקים וקבלת בלוקים בזמן fork."),
    ("4", "קונצנזוס: בעיית הגנרלים הביזנטיים, PBFT, Proof of Work, Proof of Stake, קונצנזוס Nakamoto, בחירת השרשרת הארוכה ו-finality."),
    ("5", "אפיון ותיחום איומים: מערכות שומרות פרטיות, hard fork ו-soft fork, שכבה שנייה, רעיונות Lightning/Plasma ופשרות תכן."),
    ("6", "Ethereum וחוזים חכמים: חשבונות, EVM, גז, Solidity, רשת מקומית, Remix, Hardhat או Foundry, Ganache, MetaMask, ethers.js ו-Web3.js."),
    ("7", "האקוסיסטם של Ethereum: כלכלת טוקנים, DAO, גיוס המונים בסגנון ICO, ממשל, והשפעות חברתיות וכלכליות."),
    ("8", "פיתוח DApp מאובטח: לקוח DApp, בדיקות חוזים חכמים, ספריות OpenZeppelin, reentrancy, בקרת גישה ומניפולציית oracle."),
    ("9", "טוקנים: ERC-20, ERC-721, יתרות, minting, העברות, הרשאות, דפוסי שוק NFT ושיקולי ביקורת."),
    ("10", "סקיילינג ואחסון: מושגי Ethereum 2.0, finality ב-Proof of Stake, IPFS, מטא-דאטה ל-DApp ורשתות מתחרות."),
    ("11", "ארנקים ומשמורת: HD wallets, seed phrases, אחסון חם וקר, multi-signature, תרחישי שחזור וביטול הרשאות."),
    ("12", "DeFi: AMM, הלוואות, נזילות, tokenomics, התקפות flash loan, התקפות oracle, trace, ביקורת והקשחה."),
    ("13", "פרויקט סיום ומגמות: מערכת מבוזרת עובדת, דוח ביקורת, הגנה בעל פה וכיוונים עדכניים בעולם הבלוקצ'יין."),
]

TOOLS_EN = [
    "Solidity, Remix, Hardhat, Foundry, Ganache, MetaMask, ethers.js, Web3.js",
    "OpenZeppelin Contracts, Slither, Echidna, Tenderly, IPFS, Chainlink",
]

TOOLS_HE = [
    "Solidity, Remix, Hardhat, Foundry, Ganache, MetaMask, ethers.js, Web3.js",
    "OpenZeppelin Contracts, Slither, Echidna, Tenderly, IPFS, Chainlink",
]


def clear_body(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def set_rtl(paragraph, rtl):
    if not rtl:
        return
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:bidi")) is None:
        bidi = OxmlElement("w:bidi")
        jc = p_pr.find(qn("w:jc"))
        if jc is not None:
            jc.addprevious(bidi)
        else:
            p_pr.insert(0, bidi)
    for run in paragraph.runs:
        r_pr = run._r.get_or_add_rPr()
        if r_pr.find(qn("w:rtl")) is None:
            r_pr.append(OxmlElement("w:rtl"))


def para(doc, text, style=None, rtl=False, align=None, bold=False, size=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    set_rtl(p, rtl)
    return p


def heading(doc, text, rtl=False):
    align = WD_ALIGN_PARAGRAPH.JUSTIFY if rtl else WD_ALIGN_PARAGRAPH.LEFT
    try:
        return para(doc, text, style="Heading 1", rtl=rtl, align=align)
    except KeyError:
        return para(doc, text, rtl=rtl, align=align, bold=True, size=13)


def weekly_table(doc, header, rows, rtl=False):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    if rtl:
        tbl_pr = table._tbl.tblPr
        if tbl_pr.find(qn("w:bidiVisual")) is None:
            bidi = OxmlElement("w:bidiVisual")
            tbl_style = tbl_pr.find(qn("w:tblStyle"))
            if tbl_style is not None:
                tbl_style.addnext(bidi)
            else:
                tbl_pr.insert(0, bidi)
    cells = table.rows[0].cells
    for i, text in enumerate(header):
        cells[i].text = ""
        p = cells[i].paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        if rtl:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_rtl(p, True)
    for week, subject in rows:
        cells = table.add_row().cells
        cells[0].text = week
        cells[1].text = subject
        for cell in cells:
            p = cell.paragraphs[0]
            if rtl:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                set_rtl(p, True)
    return table


def compact_body(doc, size=10):
    for p in doc.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        for run in p.runs:
            if run.font.size is None or run.font.size.pt > size + 1:
                run.font.size = Pt(size)


def build_syllabus_en():
    doc = Document(os.path.join(TEMPLATE_ROOT, "syllabus_en.docx"))
    clear_body(doc)
    left = WD_ALIGN_PARAGRAPH.LEFT
    para(doc, f"{TITLE_EN} ({OFFICIAL_TITLE_HE}) - {COURSE_NO}", bold=True, size=15, align=left)
    para(doc, "Blockchain vision and practice: consensus, smart contracts, DApps, wallets, DeFi, and security auditing.", align=left)
    para(doc, f"Lecturer: {LECTURER_EN}", align=left)
    para(doc, HOURS_EN, align=left)
    para(doc, CREDITS_EN, align=left)
    para(doc, PREREQ_EN, align=left)
    para(doc, CONCURRENT_EN, align=left)

    heading(doc, "Course Objectives")
    para(doc, "The course introduces the blockchain infrastructure and develops practical ability to understand the crypto-economy, decentralized applications, and the security limits of systems built on blockchains. Students learn how blockchain networks are structured, how peer-to-peer transaction systems work without a central controller, and how smart contracts are programmed, debugged, deployed, tested, and audited.", align=left)
    para(doc, "The course combines the original practical vision of blockchain, cryptocurrency markets, tokens, ICO-style projects, and DApps with modern security-oriented engineering: consensus under adversarial failure, smart-contract vulnerability analysis, secure wallet design, DeFi mechanisms, decentralized storage, oracle dependencies, fuzzing, static analysis, and audit reporting.", align=left)

    heading(doc, "Course content")
    para(doc, "Students build a minimal decentralized system across the semester: a ledger with cryptographic block integrity, transaction validation, consensus logic, deployed Solidity contracts, a DApp client, ERC-20 or ERC-721 token flows, IPFS-backed metadata, wallet and multi-signature controls, and a DeFi or token-funding mechanism. Every layer includes a demonstrated attack and mitigation.", align=left)

    heading(doc, "Student duties and grade components")
    para(doc, "Students submit exercises and build a final project. The project includes a requirements document, a working DApp or decentralized-system implementation, and a structured security audit report. Grade components: project 70%, exercises 30%. Passing condition: above 60 in the project and above 60 in the weighted final grade.", align=left)

    heading(doc, "Course of lessons")
    para(doc, "Teaching methods: lectures, live coding, guided practice, hands-on smart-contract development, local-chain deployment, exploit demonstrations, and project presentations. Guest lecturers: none.", align=left)
    para(doc, "Use of technology: Solidity, Ethereum tooling, local blockchain networks, static analysis, fuzzing, wallet libraries, decentralized storage, and AI-assisted coding with mandatory verification through tests and audits.", align=left)
    para(doc, "The order of the lessons may change according to teaching needs:", align=left)
    weekly_table(doc, ("Week", "Subject"), WEEKS_EN, rtl=False)

    heading(doc, "Textbooks and references")
    for item in BIB:
        para(doc, item, align=left)

    doc.save(os.path.join(OUT, "syllabus_en.docx"))


def build_syllabus_he():
    doc = Document(os.path.join(TEMPLATE_ROOT, "syllabus_he.docx"))
    clear_body(doc)
    right = WD_ALIGN_PARAGRAPH.JUSTIFY
    left = WD_ALIGN_PARAGRAPH.LEFT
    para(doc, f"{OFFICIAL_TITLE_HE} - {COURSE_NO}", rtl=True, align=right, bold=True, size=15)
    para(doc, TITLE_EN, align=left)
    para(doc, f"שם המרצה: {LECTURER_HE}", rtl=True, align=right)
    para(doc, "אופן הוראה: שיעור ותרגול", rtl=True, align=right)
    para(doc, HOURS_HE, rtl=True, align=right)
    para(doc, CREDITS_HE, rtl=True, align=right)
    para(doc, PREREQ_HE, rtl=True, align=right)
    para(doc, CONCURRENT_HE, rtl=True, align=right)

    heading(doc, "מטרת הקורס", rtl=True)
    para(doc, "הקורס מקנה היכרות מעמיקה עם תשתית ה-Blockchain ויכולת להבין את שוק הכלכלה החדשה והמטבעות הקריפטוגרפיים הבנויים עליה. הסטודנטים לומדים את מרכיבי רשת הבלוקצ'יין, את עקרונות הפעולה של רשתות Peer to Peer, ואת אופן הכתיבה, הדיבוג, הפריסה והביקורת של חוזים חכמים ויישומים מבוזרים.", rtl=True, align=right)
    para(doc, "הקורס משלב את החזון והפרקטיקה של בלוקצ'יין, שוק הקריפטו, טוקנים, ICO ו-DApps עם דגש עדכני על הנדסה ואבטחה: קונצנזוס תחת כשלים עוינים, ניתוח חולשות בחוזים חכמים, ארנקי HD ו-multi-signature, DeFi, IPFS, oracle dependencies, fuzzing, static analysis ודוח audit מובנה.", rtl=True, align=right)

    heading(doc, "במהלך הקורס הסטודנטים יידרשו", rtl=True)
    para(doc, "לבנות בקוד ארכיטקטורת Blockchain בסיסית המאפשרת לבצע עסקאות בצורה מבוזרת ללא שליטה של גורם מרכזי.", rtl=True, align=right)
    para(doc, "לפתח חוזה חכם, טוקן או יישום DApp שיופעל מעל רשת בלוקצ'יין מקומית או ניסיונית.", rtl=True, align=right)
    para(doc, "להגיש וליישם מסמך דרישות של פרויקט DApp או מערכת מבוזרת, ולסיים עם מערכת עובדת ודוח ביקורת אבטחה.", rtl=True, align=right)

    heading(doc, "חובות התלמידים ומרכיבי הציון", rtl=True)
    para(doc, "מרכיבי הציון: פרויקט 70%, תרגילים 30%. ציון עובר: מעל 60 בפרויקט ומעל 60 בשקלול כל מרכיבי הציון.", rtl=True, align=right)

    heading(doc, "שיטות הוראה ושימוש בטכנולוגיה", rtl=True)
    para(doc, "שיטות ההוראה כוללות הרצאות, הדגמות קוד, תרגול מונחה, פיתוח חוזים חכמים, פריסה לרשת מקומית, הדגמת חולשות והקשחה, ומצגות פרויקט.", rtl=True, align=right)
    para(doc, "שימוש בטכנולוגיה: " + "; ".join(TOOLS_HE), rtl=True, align=right)
    para(doc, "מרצים אורחים: אין.", rtl=True, align=right)

    heading(doc, "הנושאים שיילמדו לפי שבועות", rtl=True)
    para(doc, "סדר הנושאים יכול להשתנות בהתאם לשיקול דעת המרצה.", rtl=True, align=right)
    weekly_table(doc, ("שבוע", "נושאים"), WEEKS_HE, rtl=True)

    heading(doc, "רשימת מקורות", rtl=True)
    for item in BIB:
        para(doc, item, align=left)

    doc.save(os.path.join(OUT, "syllabus_he.docx"))


def build_rationale():
    doc = Document(os.path.join(TEMPLATE_ROOT, "rationale.docx"))
    clear_body(doc)
    right = WD_ALIGN_PARAGRAPH.JUSTIFY
    left = WD_ALIGN_PARAGRAPH.LEFT
    para(doc, f"מסמך רציונל לקורס {OFFICIAL_TITLE_HE}", rtl=True, align=right, bold=True, size=14)
    para(doc, TITLE_EN, align=left, bold=True)
    para(doc, "הקורס נדרש משום שבלוקצ'יין איננו רק נושא תיאורטי בקריפטוגרפיה או במערכות מבוזרות, אלא תחום הנדסי שבו החלטות תכן, מימוש ואבטחה משפיעות ישירות על נכסים דיגיטליים, אמון, ממשל ותשתיות פיננסיות. סטודנטים במדעי המחשב זקוקים להבנה מעשית של הרכיבים: רשת P2P, עסקאות, קונצנזוס, חוזים חכמים, ארנקים, טוקנים, DApps ו-DeFi.", rtl=True, align=right)
    para(doc, "הקורס משמר את הראייה הרחבה של בלוקצ'יין חזון ופרקטיקה, כולל שוק הקריפטו, ICO, Ethereum, DApps, רשתות חלופיות ויישומי עולם אמיתי, ומעדכן אותה בפרקטיקות הנדרשות כיום: בדיקות חוזים חכמים, fuzzing, static analysis, ביקורת קוד, ניתוח התקפות reentrancy, oracle ו-flash loan, ושימוש בכלים מודרניים כמו Foundry, Hardhat, Slither, Echidna, Tenderly, IPFS ו-Chainlink.", rtl=True, align=right)
    para(doc, "הקורס מתאים לשנה ג' משום שהוא מחבר בין מבני נתונים, תכנות מונחה עצמים, קריפטוגרפיה, רשתות ומערכות מבוזרות לכדי פרויקט הנדסי אחד. התוצר הסופי מחייב את הסטודנט לבנות, לבדוק, להסביר ולהגן על מערכת מבוזרת עובדת ודוח ביקורת אבטחה.", rtl=True, align=right)
    doc.save(os.path.join(OUT, "rationale.docx"))


def build_catalogue_summary():
    doc = Document(os.path.join(TEMPLATE_ROOT, "catalogue_summary.docx"))
    clear_body(doc)
    right = WD_ALIGN_PARAGRAPH.JUSTIFY
    left = WD_ALIGN_PARAGRAPH.LEFT
    para(doc, "תקצירים לידיעון", rtl=True, align=right, bold=True, size=14)
    para(doc, f"{OFFICIAL_TITLE_HE} - {COURSE_NO}", rtl=True, align=right, bold=True)
    para(doc, TITLE_EN, align=left)
    para(doc, "אופן הוראה: שיעור ותרגול", rtl=True, align=right)
    para(doc, HOURS_HE, rtl=True, align=right)
    para(doc, CREDITS_HE, rtl=True, align=right)
    para(doc, PREREQ_HE, rtl=True, align=right)
    para(doc, "הקורס מקנה היכרות מעשית עם תשתית ה-Blockchain, רשתות Peer to Peer, עסקאות, קונצנזוס, חוזים חכמים, DApps, טוקנים, ארנקים ו-DeFi. הסטודנטים בונים מערכת מבוזרת בסיסית, מפתחים חוזים חכמים ויישום DApp, משלבים טוקנים ואחסון מבוזר, ומבצעים בדיקות וביקורת אבטחה לחוזים ולמנגנוני DeFi.", rtl=True, align=right)
    para(doc, "נושאי הקורס: בלוקים, mempool, UTXO, Merkle trees ו-Bloom filters; קונצנזוס, PBFT, Proof of Work ו-Proof of Stake; Ethereum, EVM, Solidity, gas וכלי פיתוח; ERC-20 ו-ERC-721; IPFS; ארנקי HD ו-multi-signature; AMM, ICO ו-DeFi; התקפות reentrancy, oracle ו-flash loan; וכתיבת דוח audit.", rtl=True, align=right)

    para(doc, TITLE_EN, align=left, bold=True)
    para(doc, HOURS_EN, align=left)
    para(doc, CREDITS_EN, align=left)
    para(doc, PREREQ_EN, align=left)
    para(doc, "A practical course on blockchain infrastructure, peer-to-peer networks, transactions, consensus, smart contracts, DApps, tokens, wallets, and DeFi. Students build a basic decentralized system, develop smart contracts and a DApp, integrate token flows and decentralized storage, and conduct testing and security auditing of contracts and DeFi mechanisms.", align=left)
    para(doc, "Topics: blocks, mempools, UTXOs, Merkle trees and Bloom filters; consensus, PBFT, Proof of Work and Proof of Stake; Ethereum, EVM, Solidity, gas and development tooling; ERC-20 and ERC-721; IPFS; HD and multi-signature wallets; AMMs, ICO-style funding and DeFi; reentrancy, oracle and flash-loan attacks; and audit reporting.", align=left)
    doc.save(os.path.join(OUT, "catalogue_summary.docx"))


def build_committee_questionnaire():
    doc = Document(COMMITTEE_TEMPLATE)
    clear_body(doc)
    right = WD_ALIGN_PARAGRAPH.JUSTIFY
    left = WD_ALIGN_PARAGRAPH.LEFT
    para(doc, "שאלון לוועדת קריקולום", rtl=True, align=right, bold=True, size=14)
    para(doc, f"שם הקורס: {OFFICIAL_TITLE_HE}", rtl=True, align=right)
    para(doc, f"שם הקורס באנגלית: {TITLE_EN}", align=left)
    para(doc, f"מספר קורס: {COURSE_NO}", rtl=True, align=right)
    para(doc, f"מרצה: {LECTURER_HE}", rtl=True, align=right)
    para(doc, HOURS_HE, rtl=True, align=right)
    para(doc, CREDITS_HE, rtl=True, align=right)
    para(doc, PREREQ_HE, rtl=True, align=right)
    para(doc, CONCURRENT_HE, rtl=True, align=right)

    heading(doc, "רציונל לפתיחת או עדכון הקורס", rtl=True)
    para(doc, "הקורס מספק בסיס מעשי ועדכני בתחום הבלוקצ'יין, החוזים החכמים והמערכות המבוזרות. הוא מחבר בין מבני נתונים, תכנות, רשתות, קריפטוגרפיה ואבטחת תוכנה, ומתרגם אותם למערכת עובדת ולדוח ביקורת. העדכון מוסיף לכלים ולתכנים הקיימים פרקטיקות מודרניות: Solidity tooling, fuzzing, static analysis, DApp clients, IPFS, token standards, oracle dependencies ו-DeFi attacks.", rtl=True, align=right)

    heading(doc, "תרומת הקורס לתכנית הלימודים", rtl=True)
    para(doc, "הקורס מחזק את מסלול הסייבר ואת מסלול AI and Quantum Computing for Finance בכך שהוא מקנה לסטודנטים יכולת לבנות, להבין, לבדוק ולאבטח מערכות מבוזרות בעלות משמעות כלכלית וטכנולוגית. הוא מתאים במיוחד לתפקידים בתחומי פיתוח תוכנה פיננסית, אבטחת יישומים, smart-contract auditing, DevSecOps ו-Web3 engineering.", rtl=True, align=right)

    heading(doc, "תוצרי למידה מרכזיים", rtl=True)
    for item in [
        "הסבר עקרונות קונצנזוס, עסקאות, בלוקים ורשתות P2P.",
        "פיתוח ופריסה של חוזים חכמים ויישומי DApp.",
        "מימוש טוקנים, ארנקים ומנגנוני DeFi בסיסיים.",
        "זיהוי, ניצול והקשחה מול חולשות בחוזים חכמים.",
        "כתיבת דוח ביקורת אבטחה מובנה למערכת מבוזרת.",
    ]:
        para(doc, item, rtl=True, align=right)

    heading(doc, "אופן הערכה", rtl=True)
    para(doc, "פרויקט 70%, תרגילים 30%. תנאי מעבר: מעל 60 בפרויקט ומעל 60 בשקלול כל מרכיבי הציון.", rtl=True, align=right)
    compact_body(doc, size=10)
    doc.save(os.path.join(OUT, "committee_questionnaire.docx"))


if __name__ == "__main__":
    build_syllabus_en()
    build_syllabus_he()
    build_rationale()
    build_catalogue_summary()
    build_committee_questionnaire()
    print("Wrote CY5 catalogue package to", OUT)
